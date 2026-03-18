"""모니터링 보고서 파일 내보내기 (HTML, DOCX, XLSX)
=================================================
monitoring_report.py의 데이터 함수를 재사용하여 보고서를 파일로 내보냅니다.

사용법:
  python export_report.py <데이터_디렉토리> [비교_디렉토리]

예시:
  python export_report.py shucle_data/영덕관광/20260219_20260225 shucle_data/영덕관광/20260212_20260218

출력:
  shucle_report/{지역}/{기간}/report_{지역}_{기간}.html
  shucle_report/{지역}/{기간}/report_{지역}_{기간}.docx
  shucle_report/{지역}/{기간}/report_{지역}_{기간}.xlsx
"""

import os, sys, re
from datetime import datetime
from collections import OrderedDict

sys.stdout.reconfigure(encoding="utf-8", errors="replace")
sys.stderr.reconfigure(encoding="utf-8", errors="replace")

from monitoring_report import (
    load_charts, parse_dir_info, FRAMEWORK,
    get_kpi_value, get_drilldown_value, compute_change,
    should_trigger, fmt_val, status_label, _resolve_dynamic_name,
)


# ================================================================
# 보고서 데이터 수집
# ================================================================

def build_report_data(data_dir, compare_dir=None):
    """보고서에 필요한 모든 데이터를 구조화하여 반환"""
    region, period_fmt = parse_dir_info(data_dir)
    charts = load_charts(data_dir)
    prev_charts = load_charts(compare_dir) if compare_dir else None
    has_compare = prev_charts is not None and len(prev_charts) > 0
    _, prev_period = parse_dir_info(compare_dir) if compare_dir else ("", "")

    # 기간 코드 추출 (20260219_20260225)
    parts = os.path.normpath(data_dir).replace("\\", "/").split("/")
    period_code = ""
    for i, p in enumerate(parts):
        if p == "shucle_data" and i + 2 < len(parts):
            period_code = parts[i + 2]

    # ── KPI 값 수집 ──
    all_kpi_results = []
    for cat in FRAMEWORK:
        for kpi in cat["primary"]:
            value, daily, chart = get_kpi_value(charts, kpi)
            prev_value, prev_daily = None, OrderedDict()
            if has_compare:
                prev_value, prev_daily, _ = get_kpi_value(prev_charts, kpi)
            change = compute_change(value, prev_value)
            all_kpi_results.append({
                "cat": cat, "kpi": kpi,
                "value": value, "daily": daily, "chart": chart,
                "prev_value": prev_value, "prev_daily": prev_daily,
                "change": change,
            })

    # ── 1차 핵심지표 ──
    primary_rows = []
    triggered_list = []
    stable_list = []
    for r in all_kpi_results:
        cat = r["cat"]
        kpi = r["kpi"]
        value, prev_value, change = r["value"], r["prev_value"], r["change"]
        is_pct = kpi.get("is_pct", False)
        unit = kpi["unit"]

        curr_str = fmt_val(value, unit, is_pct) if value is not None else "-"
        prev_str = fmt_val(prev_value, unit, is_pct) if prev_value is not None else "-"

        if value is not None and prev_value is not None:
            diff = value - prev_value
            if is_pct:
                diff_str = f"{diff*100:+.1f}%p"
            else:
                diff_str = f"{diff:+.1f}{unit}" if isinstance(diff, float) else f"{diff:+,}{unit}"
        else:
            diff_str = "-"

        rate_str = f"{change*100:+.1f}%" if change is not None else "-"
        stat = status_label(change, kpi["name"])

        is_triggered = False
        if has_compare and change is not None:
            for trigger in kpi.get("triggers", []):
                if should_trigger(change, trigger):
                    is_triggered = True
                    triggered_list.append(r)
                    break

        # 트리거 안 됐지만 드릴다운이 있는 KPI → 안정 목록
        if has_compare and change is not None and not is_triggered:
            all_dds = []
            for trigger in kpi.get("triggers", []):
                all_dds.extend(trigger.get("drilldowns", []))
            if all_dds:
                stable_list.append(r)

        is_negative = stat in ("악화", "감소")
        is_positive = stat in ("개선", "증가")

        primary_rows.append({
            "category": cat["category"],
            "cat_id": cat["id"],
            "name": kpi["name"],
            "prev_str": prev_str,
            "curr_str": curr_str,
            "diff_str": diff_str,
            "rate_str": rate_str,
            "status": f"⚠ {stat}" if is_triggered else stat,
            "is_negative": is_negative,
            "is_positive": is_positive,
            "is_triggered": is_triggered,
        })

    # ── 드릴다운 행 생성 헬퍼 ──
    def _build_dd_rows(dd_list):
        rows = []
        seen = set()
        for dd in dd_list:
            if dd["name"] in seen:
                continue
            seen.add(dd["name"])

            dd_is_pct = dd.get("is_pct", False)
            dd_stats, _ = get_drilldown_value(charts, dd)
            curr_num = None
            if dd_stats and dd_stats["count"] > 0:
                curr_num = dd_stats["avg"]
                curr_dd_str = fmt_val(dd_stats["avg"], is_pct=dd_is_pct)
            else:
                curr_dd_str = "-"

            dd_display_name = _resolve_dynamic_name(dd, dd_stats)

            prev_num = None
            if has_compare and prev_charts:
                dd_prev_stats, _ = get_drilldown_value(prev_charts, dd)
                if dd_prev_stats and dd_prev_stats["count"] > 0:
                    prev_num = dd_prev_stats["avg"]
                    prev_dd_str = fmt_val(dd_prev_stats["avg"], is_pct=dd_is_pct)
                else:
                    prev_dd_str = "-"
            else:
                prev_dd_str = "-"

            for k in list(dd.keys()):
                if k.startswith("_"):
                    del dd[k]

            if curr_num is not None and prev_num is not None:
                dd_diff = curr_num - prev_num
                if dd_is_pct:
                    dd_diff_str = f"{dd_diff*100:+.1f}%p"
                else:
                    dd_diff_str = f"{dd_diff:+.2f}" if abs(dd_diff) < 10 else f"{dd_diff:+.1f}"
                dd_change = compute_change(curr_num, prev_num)
                dd_rate_str = f"{dd_change*100:+.1f}%" if dd_change is not None else "-"
            else:
                dd_diff_str = "-"
                dd_rate_str = "-"

            rows.append({
                "category": kpi["name"],
                "name": dd_display_name,
                "prev_str": prev_dd_str,
                "curr_str": curr_dd_str,
                "diff_str": dd_diff_str,
                "rate_str": dd_rate_str,
                "point": dd.get("reason", ""),
            })
        return rows

    # ── 2차 세부지표_변동 ──
    drilldown_sections = []
    for r in triggered_list:
        kpi = r["kpi"]
        change = r["change"]

        active_dds = []
        for trigger in kpi.get("triggers", []):
            if should_trigger(change, trigger):
                active_dds.extend(trigger.get("drilldowns", []))
        if not active_dds:
            continue

        dd_rows = _build_dd_rows(active_dds)
        direction = "▲" if change > 0 else "▼"
        drilldown_sections.append({
            "kpi_name": kpi["name"],
            "change_pct": f"{change*100:+.1f}%",
            "direction": direction,
            "rows": dd_rows,
        })

    # ── 2차 세부지표_안정 ──
    stable_sections = []
    for r in stable_list:
        kpi = r["kpi"]
        change = r["change"]

        all_dds = []
        for trigger in kpi.get("triggers", []):
            all_dds.extend(trigger.get("drilldowns", []))
        if not all_dds:
            continue

        dd_rows = _build_dd_rows(all_dds)
        direction = "▲" if change > 0 else ("▼" if change < 0 else "─")
        stable_sections.append({
            "kpi_name": kpi["name"],
            "change_pct": f"{change*100:+.1f}%",
            "direction": direction,
            "rows": dd_rows,
        })

    # ── 핵심 해석 ──
    insights = _build_insights(all_kpi_results, charts, has_compare)

    return {
        "region": region,
        "period_fmt": period_fmt,
        "period_code": period_code,
        "prev_period": prev_period,
        "chart_count": len(charts),
        "has_compare": has_compare,
        "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "primary_rows": primary_rows,
        "drilldown_sections": drilldown_sections,
        "stable_sections": stable_sections,
        "insights": insights,
    }


def _build_insights(all_kpi_results, charts, has_compare):
    """핵심 해석 생성 (generate_report의 insights 로직 재사용)"""
    kpi_map = {r["kpi"]["name"]: r for r in all_kpi_results}

    calls_r = kpi_map.get("실시간 호출 건수", {})
    completed_r = kpi_map.get("이동완료 호출 건수", {})
    passengers_r = kpi_map.get("총 탑승객 수", {})
    wait_r = kpi_map.get("평균 대기시간", {})
    travel_r = kpi_map.get("평균 이동시간", {})
    detour_r = kpi_map.get("평균 우회비율", {})
    success_r = kpi_map.get("가호출 성공률", {})
    dau_r = kpi_map.get("DAU (일간 활성 회원)", {})
    newmem_r = kpi_map.get("신규 지역 회원", {})

    insights = []

    # 수요 종합
    calls_v = calls_r.get("value")
    completed_v = completed_r.get("value")
    passengers_v = passengers_r.get("value")
    calls_ch = calls_r.get("change")
    completed_ch = completed_r.get("change")
    passengers_ch = passengers_r.get("change")
    calls_daily = calls_r.get("daily", OrderedDict())

    if calls_v and (calls_ch is not None and abs(calls_ch) >= 0.10):
        active_days = [d for d, v in calls_daily.items() if v > 0]
        prev_daily = calls_r.get("prev_daily", OrderedDict())
        prev_active_days = [d for d, v in prev_daily.items() if v > 0]
        day_info = ""
        if active_days:
            day_list = ",".join(active_days)
            day_info = f" — 분석기간 운행 {len(active_days)}일({day_list})"
            if prev_active_days:
                day_info += f", 비교기간 {len(prev_active_days)}일"
                if len(active_days) < len(prev_active_days):
                    day_info += "로 운행일수 차이 가능성"
        parts = []
        if calls_ch is not None:
            parts.append(f"호출({calls_ch*100:+.0f}%)")
        if completed_ch is not None and abs(completed_ch) >= 0.10:
            parts.append(f"이동완료({completed_ch*100:+.0f}%)")
        if passengers_ch is not None and abs(passengers_ch) >= 0.10:
            parts.append(f"탑승객({passengers_ch*100:+.0f}%)")
        direction = "감소" if calls_ch < 0 else "증가"
        insights.append(f"수요 {direction}: {', '.join(parts)} 모두 {'하락' if calls_ch < 0 else '상승'}{day_info}")

    # 이동완료율
    if calls_v and completed_v:
        rate = completed_v / calls_v * 100
        prev_calls = calls_r.get("prev_value")
        prev_completed = completed_r.get("prev_value")
        cancel_stats, _ = get_drilldown_value(charts, {"name": "", "match": {"includes": ["일별", "실시간 호출 결과"]}, "col": "호출취소"})
        fail_stats, _ = get_drilldown_value(charts, {"name": "", "match": {"includes": ["일별", "실시간 호출 결과"]}, "col": "배차실패"})
        cancel_total = cancel_stats["total"] if cancel_stats else 0
        fail_total = fail_stats["total"] if fail_stats else 0
        if prev_calls and prev_completed and prev_calls > 0:
            prev_rate = prev_completed / prev_calls * 100
            rate_diff = rate - prev_rate
            extra = ""
            if cancel_total or fail_total:
                extra = f" (호출취소 {cancel_total:.0f}건, 배차실패 {fail_total:.0f}건)"
            insights.append(f"이동완료율 {prev_rate:.1f}% → {rate:.1f}% ({rate_diff:+.1f}%p): 호출 {calls_v:.0f}건 중 {completed_v:.0f}건 완료{extra}")

    # 대기시간
    wait_v = wait_r.get("value")
    wait_ch = wait_r.get("change")
    if wait_v is not None and wait_ch is not None and abs(wait_ch) >= 0.10:
        prev_wait = wait_r.get("prev_value")
        top10_stats, top10_daily = get_drilldown_value(charts, {"name": "", "match": {"includes": ["상위10% 대기시간"]}})
        top10_max = ""
        if top10_daily:
            t_max_day = max(top10_daily, key=top10_daily.get)
            top10_max = f" — 특히 {t_max_day} 상위10% 대기시간 {top10_daily[t_max_day]:.1f}분으로 극단치 발생"
        direction = "급등" if wait_ch > 0 else "개선"
        insights.append(f"대기시간 {direction}: {fmt_val(prev_wait, '분') if prev_wait else '?'} → {fmt_val(wait_v, '분')} ({wait_ch*100:+.0f}%){top10_max}")

    # 가호출 성공률
    sr_v = success_r.get("value")
    sr_ch = success_r.get("change")
    if sr_v is not None and sr_ch is not None and abs(sr_ch) >= 0.10:
        prev_sr = success_r.get("prev_value")
        vhcall_stats, _ = get_drilldown_value(charts, {"name": "", "match": {"includes": ["가호출 수"], "excludes": ["순 회원", "성공", "결과", "성공률", "일별"]}})
        vhcall_str = f", 가호출 시도 일평균 {fmt_val(vhcall_stats['avg'])}건" if vhcall_stats and vhcall_stats["count"] > 0 else ""
        insights.append(
            f"가호출 성공률 {'급락' if sr_ch < 0 else '상승'}: "
            f"{fmt_val(prev_sr, '', True) if prev_sr is not None else '?'} → {fmt_val(sr_v, '', True)}"
            f"{vhcall_str}"
            f"{' — 가호출 시도는 있으나 실제 호출 전환율 매우 낮음' if sr_v < 0.05 else ''}"
        )

    # 성장
    dau_ch = dau_r.get("change")
    newmem_v = newmem_r.get("value")
    newmem_ch = newmem_r.get("change")
    cumul_stats, cumul_daily = get_drilldown_value(charts, {"name": "", "match": {"includes": ["누적 지역 회원"], "excludes": ["일별"]}})
    cumul_last = list(cumul_daily.values())[-1] if cumul_daily else None
    if (dau_ch is not None and abs(dau_ch) >= 0.10) or (newmem_ch is not None and abs(newmem_ch) >= 0.10):
        parts = []
        if dau_ch is not None and abs(dau_ch) >= 0.10:
            parts.append(f"DAU {dau_ch*100:+.0f}%")
        if newmem_ch is not None and abs(newmem_ch) >= 0.10:
            prev_nm = newmem_r.get("prev_value")
            nm_str = f"신규회원 {prev_nm:.0f}→{newmem_v:.0f}명({newmem_ch*100:+.0f}%)" if prev_nm is not None else f"신규회원 {newmem_v:.0f}명"
            parts.append(nm_str)
        cumul_str = f", 누적 {cumul_last:.0f}명" if cumul_last else ""
        gap = ""
        if calls_ch is not None and calls_ch < -0.10 and newmem_ch is not None and newmem_ch > 0.10:
            gap = " — 관심은 늘고 있으나 실 이용으로 연결이 부족"
        elif newmem_ch is not None and newmem_ch > 0.10:
            gap = " — 이용자 기반 확대 추세"
        insights.append(f"성장 지표 긍정적: {', '.join(parts)}{cumul_str}{gap}")

    # 이동시간/우회비율
    travel_v = travel_r.get("value")
    travel_ch = travel_r.get("change")
    detour_ch = detour_r.get("change")
    if travel_ch is not None and abs(travel_ch) >= 0.10:
        prev_travel = travel_r.get("prev_value")
        detour_info = ""
        if detour_ch is not None:
            detour_info = f" — 우회비율{'도 소폭 개선' if detour_ch < 0 else ' 유지' if abs(detour_ch) < 0.05 else '은 소폭 증가'}, 경로 효율성 {'향상' if travel_ch < 0 else '저하'}"
        direction = "개선" if travel_ch < 0 else "증가"
        insights.append(f"이동시간 {direction}: {fmt_val(prev_travel, '분') if prev_travel else '?'} → {fmt_val(travel_v, '분')} ({travel_ch*100:+.0f}%){detour_info}")

    # 공급 효율
    if passengers_v and completed_v and completed_v > 0:
        per_call = passengers_v / completed_v
        if per_call > 1.2 or per_call < 0.8:
            insights.append(f"호당 평균 탑승객 {per_call:.1f}명: {'동승 이용 활발' if per_call > 1.3 else '1인 이용 위주'}")

    return insights


# ================================================================
# HTML 내보내기
# ================================================================

def export_html(data, filepath):
    """HTML 파일로 내보내기"""
    def color_style(row):
        if row["is_negative"]:
            return 'color:#D32F2F; font-weight:bold;'
        elif row["is_positive"]:
            return 'color:#1565C0; font-weight:bold;'
        return ''

    html = f"""<!DOCTYPE html>
<html lang="ko">
<head>
<meta charset="UTF-8">
<title>셔클 DRT 모니터링 보고서 - {data['region']}</title>
<style>
  body {{ font-family: 'Malgun Gothic', '맑은 고딕', sans-serif; margin: 40px; color: #333; }}
  h1 {{ font-size: 20px; border-bottom: 3px solid #333; padding-bottom: 8px; }}
  .meta {{ color: #666; font-size: 13px; margin-bottom: 24px; }}
  h2 {{ font-size: 16px; margin-top: 32px; color: #222; }}
  h3 {{ font-size: 14px; margin-top: 20px; color: #444; }}
  table {{ border-collapse: collapse; width: 100%; margin: 8px 0 16px 0; font-size: 13px; }}
  th {{ background: #F5F5F5; border: 1px solid #CCC; padding: 6px 10px; text-align: center; font-weight: bold; }}
  td {{ border: 1px solid #DDD; padding: 5px 10px; }}
  td.num {{ text-align: right; }}
  td.center {{ text-align: center; }}
  .negative {{ color: #D32F2F; font-weight: bold; }}
  .positive {{ color: #1565C0; font-weight: bold; }}
  .cat-sep td {{ border-top: 2px solid #999; }}
  ul.insights {{ padding-left: 20px; }}
  ul.insights li {{ margin-bottom: 6px; line-height: 1.6; font-size: 13px; }}
  .section-divider {{ border: none; border-top: 1px solid #DDD; margin: 28px 0 8px 0; }}
  .footer {{ margin-top: 24px; font-size: 12px; color: #999; border-top: 1px solid #DDD; padding-top: 8px; }}
</style>
</head>
<body>

<h1>셔클 DRT 모니터링 보고서</h1>
<div class="meta">
  지역: <b>{data['region']}</b> | 분석기간: {data['period_fmt']}
  {'| 비교기간: ' + data['prev_period'] if data['has_compare'] else ''}<br>
  생성: {data['generated_at']} | 차트: {data['chart_count']}개 로드<br>
  지표 분석 체계: 수요/품질/공급/성장별 1차 핵심지표와 이를 세분화한 2차 세부지표
</div>

<h2>1. 1차 핵심지표 요약</h2>
<table>
<tr><th>구분</th><th>지표</th><th>비교기간</th><th>분석기간</th><th>변동값</th><th>변동률</th><th>상태</th></tr>
"""
    prev_cat = None
    for row in data["primary_rows"]:
        cat_class = ' class="cat-sep"' if prev_cat and prev_cat != row["cat_id"] else ''
        prev_cat = row["cat_id"]
        stat_class = "negative" if row["is_negative"] else ("positive" if row["is_positive"] else "")
        stat_td = f'<td class="center {stat_class}">{row["status"]}</td>'
        html += f"""<tr{cat_class}>
  <td class="center">{row['category']}</td>
  <td>{row['name']}</td>
  <td class="num">{row['prev_str']}</td>
  <td class="num">{row['curr_str']}</td>
  <td class="num">{row['diff_str']}</td>
  <td class="num">{row['rate_str']}</td>
  {stat_td}
</tr>
"""
    html += "</table>\n"

    # 2차 세부지표_변동
    html += '\n<hr class="section-divider">\n'
    if data["drilldown_sections"]:
        html += "\n<h2>2. 2차 세부지표_변동 (1차 핵심지표 ±10% 이상 변동)</h2>\n"
        for sec in data["drilldown_sections"]:
            html += f'\n<h3>{sec["direction"]} {sec["kpi_name"]} ({sec["change_pct"]})</h3>\n'
            html += '<table>\n<tr><th>구분</th><th>세부지표</th><th>비교기간</th><th>분석기간</th><th>변동값</th><th>변동률</th><th>포인트</th></tr>\n'
            for dd in sec["rows"]:
                html += f"""<tr>
  <td>{dd['category']}</td>
  <td>{dd['name']}</td>
  <td class="num">{dd['prev_str']}</td>
  <td class="num">{dd['curr_str']}</td>
  <td class="num">{dd['diff_str']}</td>
  <td class="num">{dd['rate_str']}</td>
  <td>{dd['point']}</td>
</tr>
"""
            html += "</table>\n"

    # 2차 세부지표_안정
    html += '\n<hr class="section-divider">\n'
    if data["stable_sections"]:
        html += "\n<h2>3. 2차 세부지표_안정 (10% 이내 변동)</h2>\n"
        for sec in data["stable_sections"]:
            html += f'\n<h3>{sec["direction"]} {sec["kpi_name"]} ({sec["change_pct"]})</h3>\n'
            html += '<table>\n<tr><th>구분</th><th>세부지표</th><th>비교기간</th><th>분석기간</th><th>변동값</th><th>변동률</th><th>포인트</th></tr>\n'
            for dd in sec["rows"]:
                html += f"""<tr>
  <td>{dd['category']}</td>
  <td>{dd['name']}</td>
  <td class="num">{dd['prev_str']}</td>
  <td class="num">{dd['curr_str']}</td>
  <td class="num">{dd['diff_str']}</td>
  <td class="num">{dd['rate_str']}</td>
  <td>{dd['point']}</td>
</tr>
"""
            html += "</table>\n"

    # 핵심 해석
    html += '\n<hr class="section-divider">\n'
    html += "\n<h2>4. 핵심 해석</h2>\n<ul class=\"insights\">\n"
    for ins in data["insights"]:
        html += f"  <li>{ins}</li>\n"
    html += "</ul>\n"

    html += f"""
<div class="footer">
  ※ 변동 감지 기준: 비교기간 대비 ±10% | 분석 흐름: 수요→품질→공급→성장
</div>
</body>
</html>"""

    with open(filepath, "w", encoding="utf-8") as f:
        f.write(html)
    print(f"  [HTML] {filepath}")


# ================================================================
# DOCX 내보내기
# ================================================================

def export_docx(data, filepath):
    """Word 문서로 내보내기"""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()

    # 기본 스타일
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(10)

    # 제목
    p = doc.add_heading('셔클 DRT 모니터링 보고서', level=1)
    p.runs[0].font.size = Pt(16)

    # 메타 정보
    meta = doc.add_paragraph()
    meta.add_run(f"지역: {data['region']} | 분석기간: {data['period_fmt']}")
    if data['has_compare']:
        meta.add_run(f" | 비교기간: {data['prev_period']}")
    meta.add_run(f"\n생성: {data['generated_at']} | 차트: {data['chart_count']}개 로드")
    meta.add_run(f"\n지표 분석 체계: 수요/품질/공급/성장별 1차 핵심지표와 이를 세분화한 2차 세부지표")
    meta.runs[0].font.size = Pt(9)
    meta.paragraph_format.space_after = Pt(12)

    # ── 1차 핵심지표 ──
    doc.add_heading('1. 1차 핵심지표 요약', level=2)
    headers1 = ["구분", "지표", "비교기간", "분석기간", "변동값", "변동률", "상태"]
    tbl = doc.add_table(rows=1, cols=7, style='Table Grid')
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 헤더
    for i, h in enumerate(headers1):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)

    # 데이터 행
    for row_data in data["primary_rows"]:
        row = tbl.add_row()
        vals = [row_data["category"], row_data["name"], row_data["prev_str"],
                row_data["curr_str"], row_data["diff_str"], row_data["rate_str"],
                row_data["status"]]
        for i, v in enumerate(vals):
            cell = row.cells[i]
            cell.text = v
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (0, 6) else (WD_ALIGN_PARAGRAPH.RIGHT if i >= 2 else WD_ALIGN_PARAGRAPH.LEFT)
                for run in p.runs:
                    run.font.size = Pt(9)
                    # 상태 색상
                    if i == 6:
                        if row_data["is_negative"]:
                            run.font.color.rgb = RGBColor(0xD3, 0x2F, 0x2F)
                            run.bold = True
                        elif row_data["is_positive"]:
                            run.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)
                            run.bold = True

    # 컬럼 너비
    widths = [Cm(2), Cm(4.5), Cm(2.5), Cm(2.5), Cm(2.5), Cm(2), Cm(2.5)]
    for row in tbl.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = w

    # ── 2차 세부지표_변동 ──
    if data["drilldown_sections"]:
        doc.add_heading('2. 2차 세부지표_변동 (1차 핵심지표 ±10% 이상 변동)', level=2)
        headers2 = ["구분", "세부지표", "비교기간", "분석기간", "변동값", "변동률", "포인트"]

        for sec in data["drilldown_sections"]:
            p = doc.add_paragraph()
            run = p.add_run(f'{sec["direction"]} {sec["kpi_name"]} ({sec["change_pct"]})')
            run.bold = True
            run.font.size = Pt(10)

            tbl2 = doc.add_table(rows=1, cols=7, style='Table Grid')
            tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
            for i, h in enumerate(headers2):
                cell = tbl2.rows[0].cells[i]
                cell.text = h
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.bold = True
                        run.font.size = Pt(8)

            for dd in sec["rows"]:
                row = tbl2.add_row()
                vals = [dd["category"], dd["name"], dd["prev_str"],
                        dd["curr_str"], dd["diff_str"], dd["rate_str"], dd["point"]]
                for i, v in enumerate(vals):
                    cell = row.cells[i]
                    cell.text = v
                    for p in cell.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if 2 <= i <= 5 else WD_ALIGN_PARAGRAPH.LEFT
                        for run in p.runs:
                            run.font.size = Pt(8)

            widths2 = [Cm(3), Cm(3.5), Cm(2), Cm(2), Cm(2), Cm(1.5), Cm(5)]
            for row in tbl2.rows:
                for i, w in enumerate(widths2):
                    row.cells[i].width = w

    # ── 2차 세부지표_안정 ──
    if data["stable_sections"]:
        doc.add_heading('3. 2차 세부지표_안정 (10% 이내 변동)', level=2)
        headers3 = ["구분", "세부지표", "비교기간", "분석기간", "변동값", "변동률", "포인트"]

        for sec in data["stable_sections"]:
            p = doc.add_paragraph()
            run = p.add_run(f'{sec["direction"]} {sec["kpi_name"]} ({sec["change_pct"]})')
            run.bold = True
            run.font.size = Pt(10)

            tbl3 = doc.add_table(rows=1, cols=7, style='Table Grid')
            tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
            for i, h in enumerate(headers3):
                cell = tbl3.rows[0].cells[i]
                cell.text = h
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.bold = True
                        run.font.size = Pt(8)

            for dd in sec["rows"]:
                row = tbl3.add_row()
                vals = [dd["category"], dd["name"], dd["prev_str"],
                        dd["curr_str"], dd["diff_str"], dd["rate_str"], dd["point"]]
                for i, v in enumerate(vals):
                    cell = row.cells[i]
                    cell.text = v
                    for p in cell.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if 2 <= i <= 5 else WD_ALIGN_PARAGRAPH.LEFT
                        for run in p.runs:
                            run.font.size = Pt(8)

            widths3 = [Cm(3), Cm(3.5), Cm(2), Cm(2), Cm(2), Cm(1.5), Cm(5)]
            for row in tbl3.rows:
                for i, w in enumerate(widths3):
                    row.cells[i].width = w

    # ── 핵심 해석 ──
    doc.add_heading('4. 핵심 해석', level=2)
    for ins in data["insights"]:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(ins)
        run.font.size = Pt(9)

    # 푸터
    doc.add_paragraph()
    footer = doc.add_paragraph()
    run = footer.add_run("※ 변동 감지 기준: 비교기간 대비 ±10% | 분석 흐름: 수요→품질→공급→성장")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.save(filepath)
    print(f"  [DOCX] {filepath}")


# ================================================================
# XLSX 내보내기
# ================================================================

def export_xlsx(data, filepath):
    """Excel 파일로 내보내기"""
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side

    wb = Workbook()
    ws = wb.active
    ws.title = "1차 핵심지표"

    # 스타일 정의
    header_font = Font(name='맑은 고딕', size=9, bold=True)
    header_fill = PatternFill(start_color="F5F5F5", end_color="F5F5F5", fill_type="solid")
    data_font = Font(name='맑은 고딕', size=9)
    red_font = Font(name='맑은 고딕', size=9, bold=True, color="D32F2F")
    blue_font = Font(name='맑은 고딕', size=9, bold=True, color="1565C0")
    thin_border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin'))
    center_align = Alignment(horizontal='center', vertical='center')
    right_align = Alignment(horizontal='right', vertical='center')
    left_align = Alignment(horizontal='left', vertical='center')

    # ── 제목 ──
    ws.merge_cells('A1:G1')
    ws['A1'] = f"셔클 DRT 모니터링 보고서 — {data['region']}"
    ws['A1'].font = Font(name='맑은 고딕', size=14, bold=True)

    ws.merge_cells('A2:G2')
    meta_parts = [f"분석기간: {data['period_fmt']}"]
    if data['has_compare']:
        meta_parts.append(f"비교기간: {data['prev_period']}")
    meta_parts.append(f"생성: {data['generated_at']} | 차트: {data['chart_count']}개")
    ws['A2'] = " | ".join(meta_parts)
    ws['A2'].font = Font(name='맑은 고딕', size=9, color="666666")

    ws.merge_cells('A3:G3')
    ws['A3'] = "지표 분석 체계: 수요/품질/공급/성장별 1차 핵심지표와 이를 세분화한 2차 세부지표"
    ws['A3'].font = Font(name='맑은 고딕', size=9, color="666666")

    # ── 1차 핵심지표 ──
    ws.merge_cells('A5:G5')
    ws['A5'] = "1. 1차 핵심지표 요약"
    ws['A5'].font = Font(name='맑은 고딕', size=11, bold=True)

    headers1 = ["구분", "지표", "비교기간", "분석기간", "변동값", "변동률", "상태"]
    for i, h in enumerate(headers1):
        cell = ws.cell(row=6, column=i+1, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = center_align
        cell.border = thin_border

    row_num = 7
    for row_data in data["primary_rows"]:
        vals = [row_data["category"], row_data["name"], row_data["prev_str"],
                row_data["curr_str"], row_data["diff_str"], row_data["rate_str"],
                row_data["status"]]
        for i, v in enumerate(vals):
            cell = ws.cell(row=row_num, column=i+1, value=v)
            cell.border = thin_border
            cell.font = data_font
            if i == 0:
                cell.alignment = center_align
            elif i >= 2 and i <= 5:
                cell.alignment = right_align
            elif i == 6:
                cell.alignment = center_align
                if row_data["is_negative"]:
                    cell.font = red_font
                elif row_data["is_positive"]:
                    cell.font = blue_font
            else:
                cell.alignment = left_align
        row_num += 1

    # 컬럼 너비
    col_widths = [8, 24, 14, 14, 14, 10, 12]
    for i, w in enumerate(col_widths):
        ws.column_dimensions[chr(65+i)].width = w

    # ── 2차 세부지표_변동 시트 ──
    if data["drilldown_sections"]:
        ws2 = wb.create_sheet("2차 세부지표_변동")
        headers2 = ["구분", "세부지표", "비교기간", "분석기간", "변동값", "변동률", "포인트"]

        ws2.merge_cells('A1:G1')
        ws2['A1'] = "2. 2차 세부지표_변동 (1차 핵심지표 ±10% 이상 변동)"
        ws2['A1'].font = Font(name='맑은 고딕', size=11, bold=True)

        row_num = 3
        for sec in data["drilldown_sections"]:
            ws2.merge_cells(f'A{row_num}:G{row_num}')
            ws2.cell(row=row_num, column=1,
                     value=f'{sec["direction"]} {sec["kpi_name"]} ({sec["change_pct"]})')
            ws2.cell(row=row_num, column=1).font = Font(name='맑은 고딕', size=10, bold=True)
            row_num += 1

            for i, h in enumerate(headers2):
                cell = ws2.cell(row=row_num, column=i+1, value=h)
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = center_align
                cell.border = thin_border
            row_num += 1

            for dd in sec["rows"]:
                vals = [dd["category"], dd["name"], dd["prev_str"],
                        dd["curr_str"], dd["diff_str"], dd["rate_str"], dd["point"]]
                for i, v in enumerate(vals):
                    cell = ws2.cell(row=row_num, column=i+1, value=v)
                    cell.border = thin_border
                    cell.font = data_font
                    cell.alignment = right_align if 2 <= i <= 5 else left_align
                row_num += 1
            row_num += 1  # 섹션 간 빈 행

        col_widths2 = [22, 22, 14, 14, 12, 10, 38]
        for i, w in enumerate(col_widths2):
            ws2.column_dimensions[chr(65+i)].width = w

    # ── 2차 세부지표_안정 시트 ──
    if data["stable_sections"]:
        ws2s = wb.create_sheet("2차 세부지표_안정")
        headers2s = ["구분", "세부지표", "비교기간", "분석기간", "변동값", "변동률", "포인트"]

        ws2s.merge_cells('A1:G1')
        ws2s['A1'] = "3. 2차 세부지표_안정 (10% 이내 변동)"
        ws2s['A1'].font = Font(name='맑은 고딕', size=11, bold=True)

        row_num = 3
        for sec in data["stable_sections"]:
            ws2s.merge_cells(f'A{row_num}:G{row_num}')
            ws2s.cell(row=row_num, column=1,
                      value=f'{sec["direction"]} {sec["kpi_name"]} ({sec["change_pct"]})')
            ws2s.cell(row=row_num, column=1).font = Font(name='맑은 고딕', size=10, bold=True)
            row_num += 1

            for i, h in enumerate(headers2s):
                cell = ws2s.cell(row=row_num, column=i+1, value=h)
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = center_align
                cell.border = thin_border
            row_num += 1

            for dd in sec["rows"]:
                vals = [dd["category"], dd["name"], dd["prev_str"],
                        dd["curr_str"], dd["diff_str"], dd["rate_str"], dd["point"]]
                for i, v in enumerate(vals):
                    cell = ws2s.cell(row=row_num, column=i+1, value=v)
                    cell.border = thin_border
                    cell.font = data_font
                    cell.alignment = right_align if 2 <= i <= 5 else left_align
                row_num += 1
            row_num += 1

        col_widths2s = [22, 22, 14, 14, 12, 10, 38]
        for i, w in enumerate(col_widths2s):
            ws2s.column_dimensions[chr(65+i)].width = w

    # ── 핵심 해석 시트 ──
    ws3 = wb.create_sheet("핵심 해석")
    ws3.merge_cells('A1:B1')
    ws3['A1'] = "4. 핵심 해석"
    ws3['A1'].font = Font(name='맑은 고딕', size=11, bold=True)

    for i, ins in enumerate(data["insights"]):
        cell = ws3.cell(row=i+3, column=1, value=f"• {ins}")
        cell.font = data_font
    ws3.column_dimensions['A'].width = 100

    wb.save(filepath)
    print(f"  [XLSX] {filepath}")


# ================================================================
# 메인
# ================================================================

def main():
    if len(sys.argv) < 2:
        print("사용법: python export_report.py <데이터_디렉토리> [비교_디렉토리]")
        print("예시:   python export_report.py shucle_data/영덕관광/20260219_20260225 shucle_data/영덕관광/20260212_20260218")
        return

    data_dir = sys.argv[1]
    compare_dir = sys.argv[2] if len(sys.argv) >= 3 else None

    if not os.path.isdir(data_dir):
        print(f"[오류] 디렉토리가 존재하지 않습니다: {data_dir}")
        return
    if compare_dir and not os.path.isdir(compare_dir):
        print(f"[오류] 비교 디렉토리가 존재하지 않습니다: {compare_dir}")
        return

    # 데이터 수집
    print("보고서 데이터 수집 중...")
    data = build_report_data(data_dir, compare_dir)

    # 저장 경로 생성: shucle_report/{지역}/{기간}/
    base_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "shucle_report")
    save_dir = os.path.join(base_dir, data["region"], data["period_code"])
    os.makedirs(save_dir, exist_ok=True)

    prefix = f"report_{data['region']}_{data['period_code']}"

    print(f"\n파일 내보내기: {save_dir}")
    export_html(data, os.path.join(save_dir, f"{prefix}.html"))
    export_docx(data, os.path.join(save_dir, f"{prefix}.docx"))
    export_xlsx(data, os.path.join(save_dir, f"{prefix}.xlsx"))
    print("\n완료!")


if __name__ == "__main__":
    main()
